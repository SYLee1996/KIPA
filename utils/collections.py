import os
import re

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from openai import OpenAI
from llama_index.core import Settings, VectorStoreIndex, StorageContext, QueryBundle
from llama_index.vector_stores.milvus import MilvusVectorStore
from utils.openai_like import OpenAILikeEmbedding


# vLLM 서버 설정
client = OpenAI(api_key="not-needed", base_url="http://localhost:8000/v1")


"""
===================================================================================================
UTILS - GENERAL
===================================================================================================
"""
def save_docx(main_text_genetation, main_dict, prior_list, unique_titles, output_docx_path):
    doc = Document()
    # Add the content to the Word document
    for key, value in main_text_genetation.items():
        doc.add_heading(key, level=1)
        
        if key == '선행연구 현황 및 선행연구와 본연구의 차별성':
            doc = create_research_table(doc, prior_list, [main_dict], doc_k=len(unique_titles))
        else:   
            doc.add_paragraph(value)
            
    # Save the document
    doc.save(output_docx_path)


# DOCX 파일 처리 함수
def process_docx(docx_path, section_mapping, stop_keywords):
    doc = Document(docx_path)
    ret_data_dict = {}
    current_key = None

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        text_cleaned = re.sub(r"^[^가-힣a-zA-Z]+", "", text)

        if not text_cleaned:
            continue

        if any(keyword in text_cleaned for keyword in stop_keywords):
            break

        normalized_key = None
        for key in section_mapping:
            if key in text_cleaned:
                normalized_key = section_mapping[key]
                break

        if normalized_key:
            current_key = normalized_key
            if current_key not in ret_data_dict:
                ret_data_dict[current_key] = ""
        elif current_key:
            ret_data_dict[current_key] += " " + text_cleaned

    return ret_data_dict


def add_section(doc, heading, bullet, indent_level=0):
    """Add section headers with specific bullet points and indentation"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{bullet} {heading}")
    run.bold = True
    run.font.size = Pt(12)
    if indent_level > 0:
        paragraph.paragraph_format.left_indent = Pt(indent_level * 10)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)


def add_paragraph(doc, text, bullet=None, indent_level=0):
    """Add paragraphs with optional bullet points and indentation"""
    paragraph = doc.add_paragraph()
    if bullet:
        run = paragraph.add_run(f"{bullet} {text}")
    else:
        run = paragraph.add_run(text)
    run.font.size = Pt(11)
    if indent_level > 0:
        paragraph.paragraph_format.left_indent = Pt(indent_level * 10)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = Pt(18)


def create_research_table(doc, prior_list, main_dict):
    """Add a research table to the document"""
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[1].merge(header_cells[2])
    header_cells[1].merge(header_cells[3])
    header_cells[0].text = "구 분"
    header_cells[1].text = "선행연구와의 차별성"

    for cell in header_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub_header_cells = table.rows[1].cells
    sub_header_cells[1].text = "연구목적"
    sub_header_cells[2].text = "연구방법"
    sub_header_cells[3].text = "주요연구내용"

    for cell in sub_header_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for prior in prior_list:
        row_cells = table.add_row().cells
        row_cells[0].text = prior["구분"]
        row_cells[1].text = prior["연구목적"]
        row_cells[2].text = prior["연구방법"]
        row_cells[3].text = prior["주요연구내용"]

    row_cells = table.add_row().cells
    row_cells[0].text = main_dict["구분"]
    row_cells[1].text = main_dict["연구목적"]
    row_cells[2].text = main_dict["연구방법"]
    row_cells[3].text = main_dict["주요연구내용"]


def create_word_document(main_text_genetation, prior_list, main_dict):
    """Create a Word document using the provided data"""
    doc = Document()

    # Add sections in the specified order
    for key, value in main_text_genetation.items():
        if key != '선행연구 현황 및 선행연구와 본연구의 차별성' and isinstance(value, list):  # Handle lists
            add_section(doc, key, bullet="\u25A3")
            for item in value:
                add_paragraph(doc, item, bullet="\u25CB", indent_level=1)
        
        elif key == '선행연구 현황 및 선행연구와 본연구의 차별성':  # Handle strings
            # Add the research table
            add_section(doc, "선행연구 현황 및 선행연구와 본연구의 차별성", bullet="\u25A3")
            create_research_table(doc, prior_list, main_dict)
        
        elif isinstance(value, str):  # Handle strings
            add_section(doc, key, bullet="\u25A3")
            add_paragraph(doc, value, indent_level=1)
    return doc


def read_data(f_path=None):
    if f_path is None:
        raise Exception('File path is None')
    
    _, file_ext = f_path.rsplit('.', 1)

    ret_data_list = []
    if file_ext == 'docx':
        data = Document(f'data/{f_path}')
        for paragraph in data.paragraphs:
            ret_data_list.append(paragraph.text)
    elif file_ext == 'txt':
        with open(f'data/{f_path}', 'r') as file:
            lines = file.readlines()
            for line in lines:
                line = line.replace('\n', '')
                ret_data_list.append(line)
    else:
        raise Exception(f'File Extension is not supported: {file_ext}')
    
    return ret_data_list


def get_file_paths(nodes_title):
    file_path_list = []
    file_name_list = os.listdir('data')

    for ret_name in nodes_title:
        for fname in file_name_list:
            if fname.find(f'요약_{ret_name}.txt') >= 0:
                file_path_list.append(fname)

    return list(set(file_path_list))


"""
===================================================================================================
UTILS - RAG
===================================================================================================
"""
def perform_rag(query, top_k, db_path):
    enb_model_id = 'Qwen/Qwen2.5-7B-Instruct'
    Settings.embed_model = OpenAILikeEmbedding(model=enb_model_id,
                                               api_base='http://localhost:9000/v1',
                                               api_key='na',
                                               embed_batch_size=2048)

    vector_store = MilvusVectorStore(uri=db_path,
                                     dim=3584,
                                     overwrite=False)

    # storage_context = StorageContext.from_defaults(vector_store=vector_store)
    index = VectorStoreIndex.from_vector_store(vector_store=vector_store)
    retriever = index.as_retriever(similarity_top_k=top_k, score_threshold=0.8)

    nodes = retriever.retrieve(QueryBundle(query))
    
    score_threshold = 0.8
    filtered_nodes = [node for node in nodes if node.score >= score_threshold]

    return filtered_nodes


def print_nodes(nodes):
    for i, node_with_score in enumerate(nodes):
        node = node_with_score.node
        score = node_with_score.score
        print(f"Node {i+1}:")
        print("-" * 80)
        print(f"Node ID: {node.id_}")
        print(f"Title: {node.metadata.get('title', 'No title available')}")
        print(f"Text Snippet: {node.text[:]}...")
        print(f"Score: {score}")
        print("=" * 80)


"""
===================================================================================================
UTILS - PROMPT/GEN
===================================================================================================
"""
def generate_prompt(doc_data_dict, ref_data, top_k, purpose):
    if purpose == "관련 정책현안 및 연구의 필요성":
        extracted_data = "\n".join(
            [f"""제목: {data.metadata['title']}\n내용: '{data.text}' """ for data in ref_data[:top_k]]
        )
        research_title = doc_data_dict.get("과제명", "")
        research_purpose = doc_data_dict.get("연구 목적", "")
        prompt = f"""
            - 주요 데이터: 
            {extracted_data}
            
            - 과제명:
            {research_title}
            
            - 연구 목적:
            {research_purpose}
            
            위 내용을 바탕으로 연구의 '{purpose}'을 작성한다..
            
            **작성 규칙**:
            0. 무조건 줄글로 작성하며, 어떠한 목차나 제목도 쓰지 않는다.
            1. 어떤 경우에도 글의 시작 또는 본문에 '과제명' 등의 제목을 작성하지 않는다.
            2. 반드시 두 개의 문단으로 작성한다.
            3. 첫 번째 문단에서는 '과제명', '연구 목적' 관련하여 연구의 필요성을 기술하며, 기존 문제점과 해결의 중요성을 서술한다.
            4. 두 번째 문단에서는 연구가 기여할 수 있는 구체적인 방안과 사회적, 기술적 기대 효과를 설명한다.
            5. 줄글 형식으로 작성하며, 불릿 포인트나 목록 형식을 사용하지 않는다.
            6. 입력 데이터의 맥락과 연구 목적과 논리적으로 연결된 내용을 작성해야 한다.
            
            출력 형식 예시:
            첫 번째 문단은 기존 문제점을 중심으로 연구의 필요성을 서술한다. 
            두 번째 문단은 연구가 가져올 기여와 구체적인 기대 효과를 강조한다.
        """
    
    elif purpose == "연구 목적":
        extracted_data = "\n".join(
            [f"""제목: {data.metadata['title']}\n내용: '{data.text}' """ for data in ref_data[:top_k]]
        )
        research_title = doc_data_dict.get("과제명", "")
        research_purpose = doc_data_dict.get("연구 목적", "")
        research_necessity = doc_data_dict.get("관련 정책현안 및 연구의 필요성", "")
        prompt = f"""
            아래는 연구와 관련된 입력 데이터이다:
            - 주요 데이터: 
            {extracted_data}
            
            - 과제명:
            {research_title}
            
            - 연구 목적:
            {research_purpose}
            
            - 연구 필요성:
            {research_necessity}
            
            위 내용을 바탕으로 연구의 '{purpose}'을 작성한다.
            
            **작성 규칙**:
            0. 무조건 줄글로 작성하며, 어떠한 목차나 제목도 쓰지 않는다.
            1. 어떤 경우에도 글의 시작 또는 본문에 '과제명' 등의 제목을 작성하지 않는다.
            2. 반드시 두 개의 문단으로 작성한다.
            3. 각 문단은 구체적이고 상세한 내용을 포함하며, 연구의 목적과 중요성을 명확히 드러낸다.
            4. 첫 번째 문단에서는 '과제명', '연구 목적', '연구 필요성' 관련하여 제안할 연구의 목적과 목표를 서술하고, 두 번째 문단에서는 이를 달성하기 위한 필요성과 배경을 설명한다.
            5. 입력된 연구 필요성과 주요 데이터를 논리적으로 통합하여 작성한다.
            6. 불릿 포인트나 목록 형식이 아닌 줄글로 작성한다.
            
            예시 형식:
            첫 번째 문단에서는 연구의 목적과 이를 통해 기대되는 효과를 설명한다. 
            두 번째 문단에서는 연구가 필요한 이유와 관련된 맥락을 보다 자세히 부연 설명한다.
        """
    
    elif purpose == "주요 연구내용":
        extracted_data = "\n".join([data for data in ref_data])
        prompt = f"""
            아래의 입력 데이터를 기반으로 연구의 '주요 연구내용'을 작성한다.
            - 선행 연구들의 주요 연구 내용: {extracted_data}
            - 해당 연구의 목적: {doc_data_dict.get('연구 목적', '')}
            
            **작성 규칙**:
            0. 어떤 경우에도 '주요 연구 내용' 또는 유사한 제목을 포함하지 않는다.
            1. 글의 시작 또는 본문에 '주요 연구 내용' 등의 제목을 작성하지 않는다.
            2. 연구의 구체적인 내용을 목차 형식으로 작성하여 5개의 목차를 사용하며, 각 목차는 최대 3개의 불릿 포인트로 구성한다.
            3. 각 목차는 세부적이고 명확한 설명을 포함하며, 연구의 목적과 관련된 사항을 강조한다.
            4. 결과물은 매번 일관된 형식과 구체적인 내용을 유지해야 한다.
            5. 목차별로 주요 내용이 논리적으로 연결되도록 작성한다.
            6. 목차를 표현하기위해 '#' 또는 '*' 기호를 절대 사용하지 않는다.
            7. 출력 형식을 지킨다.
            
            출력 형식 예시:
            1. 목차1 제목
                1.1. 중목차1 제목 
                    • 세부 내용1
                    • 세부 내용2
                1.2. 중목차2 제목 
                    • 세부 내용1
                    • 세부 내용2     
            
            2. 목차2 제목
                2.1. 중목차1 제목 
                    • 세부 내용1
                    • 세부 내용2
                2.2. 중목차2 제목 
                    • 세부 내용1
                    • 세부 내용2 
            ...
        """

    # 5. 연구의 목적과 주요 연구내용을 바탕으로 구체적인 데이터 수집, 분석, 및 사례 연구 방법론을 논리적으로 연결한다.
    elif purpose == "연구 추진방법":
        extracted_data = "\n".join([data for data in ref_data])
        prompt = f"""
            아래의 입력 데이터를 기반으로 연구의 '연구 추진방법'을 작성한다.
            - 주요 연구내용: {extracted_data}
            - 해당 연구의 목적: {doc_data_dict.get('연구 목적', '')}
            
            작성 규칙:
            0. 어떤 경우에도 '연구 추진 방법' 또는 유사한 제목을 포함하지 않는다.
            1. 글의 시작 또는 본문에 '연구 추진 방법' 등의 제목을 작성하지 않는다.
            2. 연구의 구체적인 내용을 목차 형식으로 작성하여 5개의 목차를 사용하며, 각 목차는 최대 3개의 불릿 포인트로 구성한다.
            3. 각 목차는 세부적이고 명확한 설명을 포함하며, 연구의 목적과 관련된 사항을 강조한다.
            4. 결과물은 매번 일관된 형식과 세부 내용을 유지해야 한다.
            5. 목차를 표현하기위해 '#' 또는 '*' 기호를 절대 사용하지 않는다.
            6. 출력 형식을 지킨다.
            
            출력 형식 예시:
            1. 목차1 제목
                1.1. 중목차1 제목 
                    • 세부 내용1
                    • 세부 내용2
                1.2. 중목차2 제목 
                    • 세부 내용1
                    • 세부 내용2     

            2. 목차2 제목
                2.1. 중목차1 제목 
                    • 세부 내용1
                    • 세부 내용2
                2.2. 중목차2 제목 
                    • 세부 내용1
                    • 세부 내용2 
            ...
        """

    elif purpose == "기대효과":
        prompt = f""" '{doc_data_dict.get('연구 목적', '')} 목적과 연구내용의 '{ref_data}' 내용을 바탕으로 '{purpose}을 '세 문단'으로 작성 해줘. (MUST DO : '세 문단'으로 정리, 줄글로 작성, 불릿 포인트로 정리하지 않기)"""
    
    return prompt


MODE_EXM_OUT_1 = 1
MODE_EXM_OUT_2 = 2
def examine_output_llm(answer, mode):
    if mode == MODE_EXM_OUT_1:
        prompt = f"""
            아래는 생성된 출력 내용을 검토하고, 특정 형식 준수 여부를 판단한 뒤 형식에 맞추어 반환하는 역할을 수행해야 한다.

            입력 내용:
            {answer}

            검토 기준:
            0. 출력 시 '#' 기호나 '##' 또는 '###', '####' 기호를 사용하지 않는다.
            1. 목차 출력 시 숫자와 • 기호만 사용한다.
            2. 출력이 아래의 예시 형식과 동일한 구조를 유지해야 한다. (목차 번호, 중목차 번호, '•' 기호 사용)
            
            출력 형식 예시:
            1. 목차1 제목
                1.1. 중목차1 제목 
                    • 세부 내용1
                    • 세부 내용2
                1.2. 중목차2 제목 
                    • 세부 내용1
                    • 세부 내용2

            2. 목차2 제목
                2.1. 중목차1 제목 
                    • 세부 내용1
                    • 세부 내용2
                2.2. 중목차2 제목 
                    • 세부 내용1
                    • 세부 내용2

            3. 목차3 제목
                3.1. 중목차1 제목 
                    • 세부 내용1
                    • 세부 내용2
                3.2. 중목차2 제목 
                    • 세부 내용1
                    • 세부 내용2

            **규칙**:
            - 입력 내용이 위 예시 형식과 동일한 구조와 형식을 준수하고 '#' 또는 '*' 기호를 사용하지 않았다면, 입력 내용을 그대로 반환한다.
            - 형식에 어긋나거나 금지된 기호가 있을 경우, 가능한 원본 의미를 유지하며 예시 형식에 맞게 수정한 최종 결과만 반환한다.
            - 형식 점검 및 수정 과정에 대해 언급하지 않는다. "입력된 내용을 검토한 결과"나 "이를 수정하여" 등과 같은 표현을 사용하지 않는다.
            - 즉, 어떠한 과정 설명 없이 최종 결과만 직접 반환한다. 수정 사실에 대한 메타 설명 없이, 오직 예시 형식을 준수하는 최종 결과만 출력한다.
        """
    elif mode == MODE_EXM_OUT_2:
        prompt = f"""
        {answer}
        위의 내용을 바탕으로 핵심 내용을 포함하여 요약정리해.

        작성 규칙:
        0. 불릿 포인트로 정리한다.
        1. 불릿 포인트의 개수는 '3개'이다.
        2. 각 불릿 포인트의 내용은 '20자' 이내로 작성한다.
        3. 각 불릿 포인트의 내용은 최대한 간결하게 개조식으로 작성한다.
        4. 각 불릿 포인트의 내용을 작성할 때 ':' 같은 특수기호를 사용하지 않는다.

        출력 형식 예시:
        • 내용1
        • 내용2
        • 내용3
        """

    answer = llm_complete(prompt=prompt, max_tokens=2048, temperature=0.1, sys_mode=MODE_SYS_PRIOR_LLM_COMP, usr_mode=None)
    return answer


MODE_SYS_LLM_COMP = 'complete'
MODE_SYS_PRIOR_LLM_COMP = 'prior_complete'
MODE_USR_PURPOSE_LLM_COMP = '연구목적'
MODE_USR_METHOD_LLM_COMP = '연구방법'
MODE_USR_CONTENT_LLM_COMP = '주요연구내용'
def llm_complete(prompt, max_tokens=2048, temperature=0.3, sys_mode=None, usr_mode=None):
    if sys_mode == None:
        sys_content = "너는 출력 형식 검수자이다. 사용자로부터 입력값을 검토하고, 특정 형식에 맞는지 확인한 뒤 맞으면 그대로 반환하고, 어긋나면 수정하여 반환한다."
    else:
        sys_content = "너는 전문적인 과제 검토자이며 내용을 분석해." if sys_mode == MODE_SYS_PRIOR_LLM_COMP else "너는 전문적인 과제 기획 검토자야. 주어진 목적과 관련하여 연구 내용 정리를 수행해."

    if not (usr_mode is None):
        if usr_mode == MODE_USR_PURPOSE_LLM_COMP:
            usr_mode_str = "각 문장은 간결하되 연구의 핵심 목적을 구체적으로 담아야 함."
        elif usr_mode == MODE_USR_METHOD_LLM_COMP:
            usr_mode_str = "각 문장은 연구와 관련된 핵심 내용을 담아야 함. 불릿 포인트 없이 논리적 흐름을 유지."
        elif usr_mode == MODE_USR_CONTENT_LLM_COMP:
            usr_mode_str = "구체적인 연구 수행 방안을 포함. 간결하면서도 실질적인 연구 절차를 중심으로 서술."
        prompt = f"{prompt}: \n 반드시 줄글 형식으로 작성하며, 요약하여 1개의 문장으로 구성. {usr_mode_str} 문장의 전체 길이는 대략 50자 이내로 제한한다."
        
    response = client.chat.completions.create(
        model="MLP-KTLim/llama-3-Korean-Bllossom-8B",
        messages=[
            {"role": "system", "content": sys_content},
            {"role": "user", "content": prompt}
        ],
        max_tokens=max_tokens,
        temperature=temperature
    )
    return response.choices[0].message.content.strip()


def process_gen(max_tokens, temperature, sys_mode, **prompt_kwargs):
    print("--------" * 20)
    prompt = generate_prompt(**prompt_kwargs)
    answer = llm_complete(prompt=prompt, max_tokens=max_tokens, temperature=temperature, sys_mode=sys_mode, usr_mode=None)
    print(prompt_kwargs['purpose'])
    print(answer)
    return answer


def prior_process_gen(raw_data_dict, sys_mode):
    print("--------" * 20)
    prior_list = []
    prior_research_list = []
    for i, (title, text) in enumerate(raw_data_dict.items()):
        prior_dict = {}
        prior_dict['구분'] = f"주요선행연구{i+1}\n\n" + f"'{title}'"
    
        for purpose in ['연구목적', '연구방법', '주요연구내용']:
            # prior_prompt = generate_prior_prompt(text_list=text, purpose=purpose3)
            prompt = f"{text} : 위 내용을 바탕으로 '{purpose}'를 '한 문단'으로 핵심 내용을 포함하여 정리해(줄글로). (MUST DO : '한 문단'으로 정리, 불릿 포인트로 정리하지 않기, 50자 이내로 작성, 최대 1문장으로 작성)"
            prior_dict[purpose] = llm_complete(prompt=prompt, max_tokens=512, temperature=0.3, sys_mode=sys_mode, usr_mode=None)
        
        print('\n\n')
        print('구분 : ', prior_dict['구분'])
        print('연구목적 : ', prior_dict['연구목적'])
        print('연구방법 : ', prior_dict['연구방법'])
        print('주요연구내용 : ', prior_dict['주요연구내용'])
        # prior_list.append(prior_dict)
        prior_research_list.append(prior_dict['주요연구내용'])

        for purpose in ['연구목적', '연구방법', '주요연구내용']:
            prompt = f"""
            {prior_dict[purpose]}
            위의 내용을 바탕으로 '{purpose}'의 핵심 내용을 포함하여 요약정리해.

            작성 규칙:
            0. 불릿 포인트로 정리한다.
            1. 불릿 포인트는 '3개'만 만든다.
            2. 각 불릿 포인트의 내용은 '20자' 이내로 작성한다.
            3. 각 불릿 포인트의 내용은 최대한 간결하게 개조식으로 작성한다.
            4. 각 불릿 포인트의 내용을 작성할 때 ':' 같은 특수기호를 사용하지 않는다.

            출력 형식 예시:
            • 내용1
            • 내용2
            • 내용3
            """
            prior_dict[purpose] = llm_complete(prompt=prompt, max_tokens=512, temperature=0.1, sys_mode=sys_mode, usr_mode=None)

        prior_list.append(prior_dict)

    return prior_list, prior_research_list